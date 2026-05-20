"""多格式文档解析器

支持 PDF、DOCX、Markdown 格式的统一解析，输出标准化的 ParsedDocument。

使用方式：
    parser = DocumentParser()
    doc = parser.parse('技术方案.pdf', section_template='numeric')
    print(f'文本：{len(doc.lines)} 行')
    print(f'章节：{len(doc.sections)} 个')
    print(f'表格：{len(doc.tables)} 个')
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass, field

from src.section_engine import SectionInferEngine, Section
from src.table_extractor import TableExtractor, Table
from src.source_locator import build_page_map

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档

    Attributes:
        text: 完整文本内容
        lines: 按行分割的文本列表
        sections: 章节列表
        tables: 表格列表
        file_type: 文件类型（pdf/docx/markdown）
        page_map: 行号 → 页码映射（PDF 有效，其他格式为空字典）
    """
    text: str
    lines: List[str]
    sections: List[Section]
    tables: List[Table]
    file_type: str = ""
    page_map: Dict[int, int] = field(default_factory=dict)

    @property
    def line_count(self) -> int:
        """行数"""
        return len(self.lines)

    @property
    def section_count(self) -> int:
        """章节数"""
        return len(self.sections)

    @property
    def table_count(self) -> int:
        """表格数"""
        return len(self.tables)

    def __repr__(self) -> str:
        return (
            f"ParsedDocument({self.file_type}: "
            f"{self.line_count}行, {self.section_count}章节, {self.table_count}表格)"
        )


class DocumentParser:
    """多格式文档解析器

    根据文件扩展名自动选择解析器，统一输出 ParsedDocument。

    支持格式：
    - PDF: PyPDF2（文本）+ pdfplumber（表格）
    - DOCX: python-docx
    - Markdown: 直接读取

    使用示例：
        parser = DocumentParser()
        doc = parser.parse('技术方案.pdf', section_template='numeric')
    """

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.md': 'markdown',
        '.markdown': 'markdown',
    }

    def parse(self, file_path: str, section_template: str = "numeric") -> ParsedDocument:
        """解析文档

        根据文件扩展名自动选择解析器。

        Args:
            file_path: 文件路径
            section_template: 章节识别模板（numeric/chinese/markdown/mixed）

        Returns:
            解析后的文档

        Raises:
            ValueError: 文件格式不支持
            FileNotFoundError: 文件不存在
        """
        import os
        if not os.path.exists(file_path):
            raise FileNotFoundError(f'文件不存在：{file_path}')

        ext = Path(file_path).suffix.lower()
        file_type = self.SUPPORTED_FORMATS.get(ext)

        if not file_type:
            supported = list(self.SUPPORTED_FORMATS.keys())
            raise ValueError(
                f'不支持的文件格式：{ext}（支持：{supported}）'
            )

        parsers = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'markdown': self._parse_markdown,
        }

        parser_func = parsers[file_type]
        return parser_func(file_path, section_template)

    def _parse_pdf(self, file_path: str, section_template: str) -> ParsedDocument:
        """PDF 解析

        使用 PyPDF2 提取文本，pdfplumber 提取表格。
        同时构建行号 → 页码映射。

        Args:
            file_path: PDF 文件路径
            section_template: 章节识别模板

        Returns:
            解析后的文档
        """
        try:
            import PyPDF2
        except ImportError:
            raise ImportError('PDF 解析需要安装 PyPDF2：pip install PyPDF2')

        # 提取文本（按页）
        pdf_pages = self._extract_pdf_text_with_pages(file_path)

        # 构建页码映射
        page_map = build_page_map(pdf_pages)

        # 合并全文
        text = '\n'.join(page_text for _, page_text in pdf_pages)

        # 乱码修复
        text = self._fix_pdf_artifacts(text)

        lines = text.split('\n')

        # 章节推断
        sections = SectionInferEngine(template=section_template).infer(lines)

        # 表格解析
        tables = TableExtractor().extract_tables(file_path)

        return ParsedDocument(
            text=text,
            lines=lines,
            sections=sections,
            tables=tables,
            file_type='pdf',
            page_map=page_map,
        )

    def _extract_pdf_text_with_pages(self, file_path: str) -> List[tuple]:
        """从 PDF 按页提取文本

        Args:
            file_path: PDF 文件路径

        Returns:
            [(页码, 页面文本), ...]

        Raises:
            Exception: PDF 解析失败
        """
        import PyPDF2

        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ''
                    pages.append((page_num, page_text))

            # 检查是否为空
            total_text = '\n'.join(pt for _, pt in pages)
            if not total_text.strip():
                logger.warning(f'PDF 文本为空（可能为扫描件）：{file_path}')

            return pages
        except Exception as e:
            raise Exception(f'PDF 解析失败：{file_path}，错误：{e}')

    def _extract_pdf_text(self, file_path: str) -> str:
        """从 PDF 提取文本

        Args:
            file_path: PDF 文件路径

        Returns:
            完整文本

        Raises:
            Exception: PDF 解析失败
        """
        import PyPDF2

        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                full_text = ''
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ''
                    full_text += page_text + '\n'

            if not full_text.strip():
                logger.warning(f'PDF 文本为空（可能为扫描件）：{file_path}')

            return full_text
        except Exception as e:
            raise Exception(f'PDF 解析失败：{file_path}，错误：{e}')

    def _fix_pdf_artifacts(self, text: str) -> str:
        """修复 PDF 常见乱码

        PDF 提取的文本中，部分中文字符会被错误提取为 Unicode 部首。

        Args:
            text: 原始文本

        Returns:
            修复后的文本
        """
        replacements = {
            '⼿': '手', '⽤': '用', '⼾': '户', '⼆': '二', '⼭': '山',
            '⼝': '口', '⼥': '女', '⼦': '子', '⽂': '文', '⽅': '方',
            '⽆': '无', '⽇': '日', '⽉': '月', '⽊': '木', '⽔': '水',
            '⽕': '火', '⼟': '土', '⾦': '金', '⽽': '而', '⼏': '几',
            '⽬': '目', '⽴': '立', '⽵': '竹', '⽶': '米',
            '⻔': '门', '⻋': '车', '⾛': '走', '⾜': '足', '⾁': '肉',
            '⾷': '食', '⾔': '言', '⻆': '角', '⻣': '骨', '⾼': '高',
        }

        for wrong, correct in replacements.items():
            if wrong in text:
                text = text.replace(wrong, correct)

        # 清理所有不可见控制字符（\x00-\x08 \x0b \x0c \x0e-\x1f \x7f \x80-\x9f \xad）
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\xad]', '', text)
        # 清理零宽字符
        text = re.sub(r'[\u200b-\u200f\u2028-\u202f\ufeff]', '', text)
        # 替换全角空格为半角
        text = text.replace('\u3000', ' ')
        # 多个空白合并
        text = re.sub(r'[ \t]+', ' ', text)
        # 清理行首行尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text

    def _parse_docx(self, file_path: str, section_template: str) -> ParsedDocument:
        """DOCX 解析

        使用 python-docx 提取段落和表格。

        Args:
            file_path: DOCX 文件路径
            section_template: 章节识别模板

        Returns:
            解析后的文档
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError('DOCX 解析需要安装 python-docx：pip install python-docx')

        doc = Document(file_path)
        lines = [p.text for p in doc.paragraphs]
        text = '\n'.join(lines)

        # 章节推断
        sections = SectionInferEngine(template=section_template).infer(lines)

        # DOCX 表格
        tables = []
        for idx, table in enumerate(doc.tables):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue
            # 过滤空行
            non_empty_rows = [row for row in rows[1:] if any(cell.strip() for cell in row)]
            if non_empty_rows:
                tables.append(Table(
                    title=f"表格{idx + 1}",
                    header=rows[0],
                    rows=non_empty_rows
                ))

        return ParsedDocument(
            text=text,
            lines=lines,
            sections=sections,
            tables=tables,
            file_type='docx'
        )

    def _parse_markdown(self, file_path: str, section_template: str) -> ParsedDocument:
        """Markdown 解析

        直接读取文本，提取 Markdown 表格。

        Args:
            file_path: Markdown 文件路径
            section_template: 章节识别模板

        Returns:
            解析后的文档
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        lines = text.split('\n')

        # 章节推断
        sections = SectionInferEngine(template=section_template).infer(lines)

        # Markdown 表格
        tables = self._extract_markdown_tables(lines)

        return ParsedDocument(
            text=text,
            lines=lines,
            sections=sections,
            tables=tables,
            file_type='markdown'
        )

    def _extract_markdown_tables(self, lines: List[str]) -> List[Table]:
        """提取 Markdown 表格

        识别格式：
        ```
        | 列1 | 列2 |
        | --- | --- |
        | 数据1 | 数据2 |
        ```

        Args:
            lines: 文档行列表

        Returns:
            表格列表
        """
        tables: List[Table] = []
        in_table = False
        current_rows: List[List[str]] = []

        for line in lines:
            stripped = line.strip()

            # 检测表格行
            if stripped.startswith('|') and stripped.endswith('|'):
                if not in_table:
                    in_table = True
                    current_rows = []

                # 解析行
                cells = [cell.strip() for cell in stripped.strip('|').split('|')]

                # 跳过分隔行（|---|---|）
                if self._is_separator_row(cells):
                    continue

                current_rows.append(cells)

            else:
                # 表格结束
                if in_table and len(current_rows) >= 2:
                    tables.append(Table(
                        header=current_rows[0],
                        rows=current_rows[1:]
                    ))
                    current_rows = []
                in_table = False

        # 处理最后一个表格
        if in_table and len(current_rows) >= 2:
            tables.append(Table(
                header=current_rows[0],
                rows=current_rows[1:]
            ))

        return tables

    def _is_separator_row(self, cells: List[str]) -> bool:
        """判断是否为分隔行（|---|---|）

        Args:
            cells: 单元格列表

        Returns:
            是否为分隔行
        """
        for cell in cells:
            cleaned = cell.replace('-', '').replace(':', '').replace(' ', '')
            if cleaned:
                return False
        return True
